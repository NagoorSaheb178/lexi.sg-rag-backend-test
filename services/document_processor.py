import os
import logging
import re
import subprocess
from typing import List, Optional
from pathlib import Path
import PyPDF2
import docx
from dataclasses import dataclass

logger = logging.getLogger(__name__)

@dataclass
class Document:
    """Represents a processed document chunk"""
    content: str
    source: str
    chunk_id: str
    metadata: dict

class DocumentProcessor:
    """Service for processing PDF and DOCX documents"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def process_documents_from_directory(self, directory: str) -> List[Document]:
        """
        Process all PDF and DOCX files in a directory
        
        Args:
            directory: Path to directory containing documents
            
        Returns:
            List of Document objects
        """
        documents = []
        
        if not os.path.exists(directory):
            logger.warning(f"Directory {directory} does not exist")
            return documents
        
        for filename in os.listdir(directory):
            file_path = os.path.join(directory, filename)
            
            if filename.lower().endswith('.pdf'):
                docs = self._process_pdf(file_path)
                documents.extend(docs)
            elif filename.lower().endswith('.docx'):
                docs = self._process_docx(file_path)
                documents.extend(docs)
            elif filename.lower().endswith('.doc'):
                docs = self._process_doc(file_path)
                documents.extend(docs)
            elif filename.lower().endswith('.txt'):
                docs = self._process_txt(file_path)
                documents.extend(docs)
            else:
                logger.info(f"Skipping unsupported file: {filename}")
        
        logger.info(f"Processed {len(documents)} document chunks from {directory}")
        return documents
    
    def _process_pdf(self, file_path: str) -> List[Document]:
        """Process a PDF file"""
        documents = []
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num} of {file_path}: {e}")
                
                if text.strip():
                    # Clean and chunk the text
                    cleaned_text = self._clean_text(text)
                    chunks = self._chunk_text(cleaned_text)
                    
                    filename = os.path.basename(file_path)
                    for i, chunk in enumerate(chunks):
                        documents.append(Document(
                            content=chunk,
                            source=filename,
                            chunk_id=f"{filename}_chunk_{i}",
                            metadata={
                                "file_type": "pdf",
                                "file_path": file_path,
                                "chunk_index": i,
                                "total_chunks": len(chunks)
                            }
                        ))
                else:
                    logger.warning(f"No text extracted from PDF: {file_path}")
                    
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
        
        return documents
    
    def _process_docx(self, file_path: str) -> List[Document]:
        """Process a DOCX file"""
        documents = []
        
        try:
            doc = docx.Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            if text.strip():
                # Clean and chunk the text
                cleaned_text = self._clean_text(text)
                chunks = self._chunk_text(cleaned_text)
                
                filename = os.path.basename(file_path)
                for i, chunk in enumerate(chunks):
                    documents.append(Document(
                        content=chunk,
                        source=filename,
                        chunk_id=f"{filename}_chunk_{i}",
                        metadata={
                            "file_type": "docx",
                            "file_path": file_path,
                            "chunk_index": i,
                            "total_chunks": len(chunks)
                        }
                    ))
            else:
                logger.warning(f"No text extracted from DOCX: {file_path}")
                
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
        
        return documents
    
    def _process_doc(self, file_path: str) -> List[Document]:
        """Process a DOC file using antiword"""
        documents = []
        
        try:
            # Use antiword to extract text from .doc files
            result = subprocess.run(['antiword', file_path], 
                                  capture_output=True, text=True)
            
            if result.returncode == 0:
                text = result.stdout
                
                if text.strip():
                    # Clean and chunk the text
                    cleaned_text = self._clean_text(text)
                    chunks = self._chunk_text(cleaned_text)
                    
                    filename = os.path.basename(file_path)
                    for i, chunk in enumerate(chunks):
                        documents.append(Document(
                            content=chunk,
                            source=filename,
                            chunk_id=f"{filename}_chunk_{i}",
                            metadata={
                                "file_type": "doc",
                                "file_path": file_path,
                                "chunk_index": i,
                                "total_chunks": len(chunks)
                            }
                        ))
                else:
                    logger.warning(f"No text extracted from DOC: {file_path}")
            else:
                logger.error(f"Antiword failed for {file_path}: {result.stderr}")
                
        except Exception as e:
            logger.error(f"Error processing DOC {file_path}: {e}")
        
        return documents
    
    def _process_txt(self, file_path: str) -> List[Document]:
        """Process a TXT file"""
        documents = []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            if text.strip():
                # Clean and chunk the text
                cleaned_text = self._clean_text(text)
                chunks = self._chunk_text(cleaned_text)
                
                filename = os.path.basename(file_path)
                for i, chunk in enumerate(chunks):
                    documents.append(Document(
                        content=chunk,
                        source=filename,
                        chunk_id=f"{filename}_chunk_{i}",
                        metadata={
                            "file_type": "txt",
                            "file_path": file_path,
                            "chunk_index": i,
                            "total_chunks": len(chunks)
                        }
                    ))
            else:
                logger.warning(f"No text found in TXT file: {file_path}")
                
        except Exception as e:
            logger.error(f"Error processing TXT {file_path}: {e}")
        
        return documents
    
    def _clean_text(self, text: str) -> str:
        """
        Clean extracted text to remove garbled characters and formatting issues
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove garbled characters and encoding issues
        text = re.sub(r'[^\x00-\x7F]+', ' ', text)  # Remove non-ASCII characters
        text = re.sub(r'[^\w\s\.\,\!\?\;\:\(\)\[\]\{\}\"\'`\-\+\=\@\#\$\%\^\&\*]', ' ', text)
        
        # Remove multiple spaces
        text = re.sub(r'\s+', ' ', text)
        
        # Remove empty lines
        text = re.sub(r'\n\s*\n', '\n', text)
        
        # Strip and ensure proper formatting
        text = text.strip()
        
        # Replace common legal document formatting issues
        text = re.sub(r'\s+([.,!?;:])', r'\1', text)  # Fix spacing before punctuation
        text = re.sub(r'([.!?])\s*([a-z])', r'\1 \2', text)  # Fix spacing after sentences
        
        return text
    
    def _chunk_text(self, text: str) -> List[str]:
        """
        Split text into chunks with overlap
        
        Args:
            text: Text to chunk
            
        Returns:
            List of text chunks
        """
        if len(text) <= self.chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            
            # If we're not at the end, try to break at a sentence boundary
            if end < len(text):
                # Look for sentence endings within the last 100 characters
                sentence_end = text.rfind('.', start, end)
                if sentence_end > start + self.chunk_size - 100:
                    end = sentence_end + 1
                else:
                    # If no sentence ending found, break at word boundary
                    space_pos = text.rfind(' ', start, end)
                    if space_pos > start:
                        end = space_pos
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # Move start position with overlap
            start = end - self.chunk_overlap
            if start >= len(text):
                break
        
        return chunks
